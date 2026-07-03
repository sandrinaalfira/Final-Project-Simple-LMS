from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

document = Document()

# Judul
heading = document.add_heading('Daftar Paket Fitur Tambahan Terpilih', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Berikut adalah 4 (empat) fitur tambahan tingkat lanjut (advanced) yang telah berhasil diimplementasikan pada aplikasi Simple LMS. Kombinasi keempat fitur ini bernilai total 54 Poin.\n')

# Membuat Tabel
table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Mengatur Header Tabel
hdr_cells = table.rows[0].cells
headers = ['No', 'Fitur Tambahan', 'Kategori Paket (Berdasarkan PDF)', 'Poin']
for i, header_text in enumerate(headers):
    hdr_cells[i].text = header_text
    # Memberi warna abu-abu pada header
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

# Data Fitur
data = [
    (1, 'Redis Caching untuk Course', 'Paket 4 (Performance)', 12),
    (2, 'API Rate Limiting berbasis Redis', 'Paket 4 (Performance)', 12),
    (3, 'Cache Invalidation Strategy', 'Paket 4 (Performance)', 12),
    (4, 'Generate Certificate Asinkron (PDF)', 'Paket 6 (Async & Notification)', 18),
]

# Memasukkan Data ke Tabel
for no, fitur, paket, poin in data:
    row_cells = table.add_row().cells
    row_cells[0].text = str(no)
    row_cells[1].text = fitur
    row_cells[2].text = paket
    row_cells[3].text = str(poin)

# Penutup Total Poin
p = document.add_paragraph('\nTotal Nilai Poin Terkumpul: ')
r = p.add_run('54 Poin')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 102, 204) # Warna biru

p2 = document.add_paragraph()
r2 = p2.add_run('(Melampaui batas minimum pengerjaan fitur tambahan yaitu 50 Poin).')
r2.italic = True

document.save('/code/Tabel_Fitur_Terpilih.docx')
print("Document saved successfully at /code/Tabel_Fitur_Terpilih.docx")
